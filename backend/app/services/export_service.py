import io
import json
from typing import Any
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.forecast import StructuredForecastOutput


def export_forecast_json(forecast: StructuredForecastOutput) -> str:
    return forecast.model_dump_json(indent=2)


def export_forecast_pdf(forecast: StructuredForecastOutput, simulation_title: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom professional styles
    primary_color = colors.HexColor("#0f172a")  # Slate 900
    secondary_color = colors.HexColor("#2563eb") # Blue 600
    accent_color = colors.HexColor("#0284c7")    # Sky 600
    text_color = colors.HexColor("#334155")      # Slate 700
    
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Heading1'],
        fontSize=26,
        leading=32,
        textColor=primary_color,
        fontName='Helvetica-Bold',
        spaceAfter=12
    )
    
    subtitle_style = ParagraphStyle(
        'ReportSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        leading=18,
        textColor=secondary_color,
        fontName='Helvetica-Bold',
        spaceAfter=24
    )

    h2_style = ParagraphStyle(
        'ReportH2',
        parent=styles['Heading2'],
        fontSize=18,
        leading=22,
        textColor=primary_color,
        fontName='Helvetica-Bold',
        spaceBefore=16,
        spaceAfter=10
    )

    h3_style = ParagraphStyle(
        'ReportH3',
        parent=styles['Heading3'],
        fontSize=13,
        leading=16,
        textColor=accent_color,
        fontName='Helvetica-Bold',
        spaceBefore=12,
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['BodyText'],
        fontSize=10.5,
        leading=15,
        textColor=text_color,
        fontName='Helvetica',
        spaceAfter=10
    )

    bullet_style = ParagraphStyle(
        'ReportBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=6
    )

    story = []

    # Title Banner
    story.append(Paragraph("CivitasAI Strategic Foresight Report", title_style))
    story.append(Paragraph(f"Simulation Analysis: {simulation_title}", subtitle_style))
    story.append(Spacer(1, 10))

    # Executive Summary
    story.append(Paragraph("Executive Summary", h2_style))
    # Replace markdown headers or bold symbols in executive summary
    exec_text = forecast.executive_summary.replace("# Executive Foresight Report: CivitasAI Civilizational Projection", "")
    exec_text = exec_text.replace("### ", "").replace("**", "")
    for para in exec_text.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), body_style))
    story.append(Spacer(1, 15))

    # Confidence Scores Table
    story.append(Paragraph("Key Prediction Confidence Scores", h2_style))
    table_data = [["Strategic Milestone / Prediction", "Confidence Level"]]
    for score in forecast.confidence_scores:
        table_data.append([score.item, f"{score.confidence_percentage}%"])

    t = Table(table_data, colWidths=[350, 150])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), primary_color),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 11),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
        ('TOPPADDING', (0,0), (-1,0), 8),
        ('ALIGN', (1,0), (1,-1), 'CENTER'),
        ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#f8fafc")),
        ('GRID', (0,0), (-1,-1), 1, colors.HexColor("#e2e8f0")),
        ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,1), (-1,-1), 10),
        ('BOTTOMPADDING', (0,1), (-1,-1), 8),
        ('TOPPADDING', (0,1), (-1,-1), 8),
    ]))
    story.append(t)
    story.append(Spacer(1, 20))
    story.append(PageBreak())

    # Milestone Forecasts
    story.append(Paragraph("Milestone Horizon Forecasts", h2_style))
    
    for yr in [2030, 2050, 2100]:
        yr_data = forecast.forecasts.get(yr)
        if not yr_data:
            continue
        
        story.append(Paragraph(f"Horizon {yr}", h2_style))
        story.append(Paragraph("<b>Technology Domain:</b>", h3_style))
        story.append(Paragraph(f"<b>AGI & Compute:</b> {yr_data.technology.agi}", body_style))
        story.append(Paragraph(f"<b>Robotics:</b> {yr_data.technology.robotics}", body_style))
        story.append(Paragraph(f"<b>Quantum & Biotech:</b> {yr_data.technology.quantum_computing} {yr_data.technology.biotechnology}", body_style))

        story.append(Paragraph("<b>Climate & Environment:</b>", h3_style))
        story.append(Paragraph(f"Projected Sea Level Rise: {yr_data.climate.sea_level_rise_cm} cm | Global Temp Change: +{yr_data.climate.global_temperature_change_c} °C | Carbon Emissions: {yr_data.climate.carbon_emissions_gt} Gt", body_style))
        story.append(Paragraph(f"Biodiversity: {yr_data.climate.biodiversity_index}", body_style))

        story.append(Paragraph("<b>Economy & Geopolitics:</b>", h3_style))
        story.append(Paragraph(f"<b>GDP Growth Rate:</b> {yr_data.economy.gdp_growth_rate}% | <b>Automation Impact:</b> {yr_data.economy.automation_impact}", body_style))
        story.append(Paragraph(f"<b>Geopolitical Alignment:</b> {yr_data.geopolitics.alliances} {yr_data.geopolitics.emerging_powers}", body_style))

        story.append(Paragraph("<b>Space & Interplanetary Exploration:</b>", h3_style))
        story.append(Paragraph(f"<b>Lunar/Mars status:</b> {yr_data.space.lunar_colonies} {yr_data.space.mars_missions}", body_style))
        story.append(Spacer(1, 15))

    story.append(PageBreak())

    # Risk and Opportunity Analysis
    story.append(Paragraph("Risk & Opportunity Assessment", title_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Major Strategic Threats & Systemic Risks", h2_style))
    for threat in forecast.risk_analysis.major_threats + forecast.risk_analysis.systemic_risks:
        story.append(Paragraph(f"• {threat}", bullet_style))
    
    story.append(Spacer(1, 10))
    story.append(Paragraph("Black Swan Potential Events", h3_style))
    for swan in forecast.risk_analysis.black_swan_events:
        story.append(Paragraph(f"• {swan}", bullet_style))

    story.append(Spacer(1, 15))
    story.append(Paragraph("Emerging Opportunities & Industries", h2_style))
    for ind in forecast.opportunity_analysis.emerging_industries:
        story.append(Paragraph(f"• **Industry:** {ind}", bullet_style))
    for opp in forecast.opportunity_analysis.investment_opportunities:
        story.append(Paragraph(f"• **Investment:** {opp}", bullet_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def export_forecast_docx(forecast: StructuredForecastOutput, simulation_title: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    title = doc.add_heading(f"CivitasAI Strategic Foresight Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Simulation Analysis: {simulation_title}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive Summary", level=1)
    exec_text = forecast.executive_summary.replace("# Executive Foresight Report: CivitasAI Civilizational Projection", "")
    exec_text = exec_text.replace("### ", "").replace("**", "")
    for para in exec_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading("Key Prediction Confidence Scores", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Strategic Milestone / Prediction"
    hdr_cells[1].text = "Confidence Level"
    
    for score in forecast.confidence_scores:
        row_cells = table.add_row().cells
        row_cells[0].text = score.item
        row_cells[1].text = f"{score.confidence_percentage}%"

    doc.add_page_break()

    doc.add_heading("Milestone Horizon Forecasts", level=1)
    for yr in [2030, 2050, 2100]:
        yr_data = forecast.forecasts.get(yr)
        if not yr_data:
            continue
        
        doc.add_heading(f"Horizon {yr}", level=2)
        p = doc.add_paragraph()
        p.add_run("Technology Domain:\n").bold = True
        p.add_run(f"AGI & Compute: {yr_data.technology.agi}\n")
        p.add_run(f"Robotics: {yr_data.technology.robotics}\n")
        
        p2 = doc.add_paragraph()
        p2.add_run("Climate & Environment:\n").bold = True
        p2.add_run(f"Sea Level Rise: {yr_data.climate.sea_level_rise_cm} cm | Global Temp Change: +{yr_data.climate.global_temperature_change_c} °C\n")
        p2.add_run(f"Biodiversity: {yr_data.climate.biodiversity_index}\n")

        p3 = doc.add_paragraph()
        p3.add_run("Economy & Geopolitics:\n").bold = True
        p3.add_run(f"GDP Growth Rate: {yr_data.economy.gdp_growth_rate}% | Automation Impact: {yr_data.economy.automation_impact}\n")
        p3.add_run(f"Geopolitics: {yr_data.geopolitics.alliances}\n")

        p4 = doc.add_paragraph()
        p4.add_run("Space Exploration:\n").bold = True
        p4.add_run(f"Lunar/Mars: {yr_data.space.lunar_colonies} {yr_data.space.mars_missions}\n")

    doc.add_page_break()
    doc.add_heading("Risk & Opportunity Assessment", level=1)
    
    doc.add_heading("Major Threats & Systemic Risks", level=2)
    for threat in forecast.risk_analysis.major_threats + forecast.risk_analysis.systemic_risks:
        doc.add_paragraph(threat, style='List Bullet')
        
    doc.add_heading("Emerging Opportunities & Future Industries", level=2)
    for ind in forecast.opportunity_analysis.emerging_industries:
        doc.add_paragraph(f"Industry: {ind}", style='List Bullet')
    for opp in forecast.opportunity_analysis.investment_opportunities:
        doc.add_paragraph(f"Investment: {opp}", style='List Bullet')

    doc.save(buffer)
    buffer.seek(0)
    return buffer
